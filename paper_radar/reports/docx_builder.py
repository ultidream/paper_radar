from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from paper_radar.models import Paper, PaperAnalysis


class DocxReportBuilder:
    def build(
        self,
        output_path: Path,
        papers: list[tuple[Paper, PaperAnalysis]],
        report_date: date,
        owner_name: str = "",
    ) -> Path:
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc = Document()
        self._set_styles(doc)

        title = doc.add_paragraph(style="Title")
        title.add_run("每日科研论文雷达")
        subtitle = doc.add_paragraph(style="Subtitle")
        subtitle.add_run(f"{report_date.isoformat()} | {owner_name or 'Research briefing'}")

        lead = doc.add_paragraph()
        lead.add_run(f"今日共发现 {len(papers)} 篇新论文。").bold = True
        lead.add_run(" 本报告按期刊新发表记录自动整理，AI 分析用于辅助筛选；重要论文建议继续阅读全文核查。")

        self._add_summary_table(doc, papers)

        for index, (paper, analysis) in enumerate(papers, start=1):
            if index > 1:
                doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
            self._add_paper_section(doc, index, paper, analysis)

        doc.save(output_path)
        return output_path

    @staticmethod
    def _set_styles(doc: Document) -> None:
        section = doc.sections[0]
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

        styles = doc.styles
        normal = styles["Normal"]
        normal.font.name = "Arial"
        normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        normal.font.size = Pt(10)

        for style_name, size, color in [
            ("Title", 22, RGBColor(31, 78, 121)),
            ("Subtitle", 10, RGBColor(89, 89, 89)),
            ("Heading 1", 14, RGBColor(31, 78, 121)),
            ("Heading 2", 11, RGBColor(55, 96, 146)),
        ]:
            style = styles[style_name]
            style.font.name = "Arial"
            style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
            style.font.size = Pt(size)
            style.font.color.rgb = color

    @staticmethod
    def _add_summary_table(doc: Document, papers: list[tuple[Paper, PaperAnalysis]]) -> None:
        doc.add_heading("今日速览", level=1)
        table = doc.add_table(rows=1, cols=5)
        table.style = "Table Grid"
        table.autofit = False
        widths = [0.35, 0.95, 1.25, 3.35, 2.35]
        headers = ["#", "发表日期", "期刊", "论文与作者", "一句话判断"]
        for col_idx, (cell, header) in enumerate(zip(table.rows[0].cells, headers)):
            DocxReportBuilder._set_cell_width(cell, widths[col_idx])
            cell.text = header
            DocxReportBuilder._shade_cell(cell, "D9EAF7")
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(9)

        for idx, (paper, analysis) in enumerate(papers, start=1):
            cells = table.add_row().cells
            values = [
                str(idx),
                paper.published_date or "Unknown",
                paper.journal,
                DocxReportBuilder._paper_summary_text(paper),
                analysis.one_sentence,
            ]
            for col_idx, (cell, value) in enumerate(zip(cells, values)):
                DocxReportBuilder._set_cell_width(cell, widths[col_idx])
                cell.text = value
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_after = Pt(0)
                    for run in paragraph.runs:
                        run.font.size = Pt(8.5 if col_idx != 3 else 8)

    def _add_paper_section(self, doc: Document, index: int, paper: Paper, analysis: PaperAnalysis) -> None:
        doc.add_heading(f"{index}. {paper.title}", level=1)

        self._add_labeled_paragraph(doc, "基本信息", self._metadata_block(paper))
        self._add_labeled_paragraph(doc, "这篇论文做了什么", analysis.one_sentence)
        self._add_labeled_paragraph(doc, "为什么重要", analysis.why_it_matters)
        self._add_labeled_paragraph(doc, "方法 / 数据 / 模型", analysis.methods_data)
        self._add_bullets(doc, "我能学习什么", analysis.what_to_learn)
        self._add_labeled_paragraph(doc, "和我的研究方向的关系", analysis.relevance)
        self._add_bullets(doc, "可关注或合作的作者", analysis.collaboration_authors)
        self._add_bullets(doc, "合作切入点", analysis.collaboration_ideas)
        self._add_labeled_paragraph(doc, "联系作者的邮件角度", analysis.outreach_angle)

        if paper.abstract:
            self._add_labeled_paragraph(doc, "摘要", paper.abstract)

    @staticmethod
    def _paper_summary_text(paper: Paper) -> str:
        lines = [paper.title]
        if paper.first_author:
            lines.append(f"第一作者：{paper.first_author}")
        if paper.first_author_affiliations:
            lines.append(f"第一作者单位：{'; '.join(paper.first_author_affiliations[:2])}")
        if paper.corresponding_authors:
            lines.append(f"通讯作者：{', '.join(paper.corresponding_authors[:4])}")
        elif paper.authors:
            lines.append("通讯作者：公开元数据未标注")
        return "\n".join(lines)

    @staticmethod
    def _metadata_block(paper: Paper) -> str:
        lines = [
            f"期刊：{paper.journal}",
            f"发表日期：{paper.published_date or 'Unknown'}",
        ]
        if paper.doi:
            lines.append(f"DOI：{paper.doi}")
        if paper.url:
            lines.append(f"链接：{paper.url}")
        if paper.first_author:
            lines.append(f"第一作者：{paper.first_author}")
        if paper.first_author_affiliations:
            lines.append(f"第一作者单位：{'; '.join(paper.first_author_affiliations[:5])}")
        elif paper.affiliations:
            lines.append(f"作者单位：{'; '.join(paper.affiliations[:5])}")
        else:
            lines.append("作者单位：公开元数据未提供")
        if paper.corresponding_authors:
            lines.append(f"通讯作者：{', '.join(paper.corresponding_authors)}")
        else:
            lines.append("通讯作者：公开元数据未标注")
        if paper.authors:
            authors = ", ".join(paper.authors[:18])
            if len(paper.authors) > 18:
                authors += " et al."
            lines.append(f"作者列表：{authors}")
        return "\n".join(lines)

    @staticmethod
    def _add_labeled_paragraph(doc: Document, label: str, text: str) -> None:
        doc.add_heading(label, level=2)
        para = doc.add_paragraph(text or "未提供。")
        para.paragraph_format.space_after = Pt(6)

    @staticmethod
    def _add_bullets(doc: Document, label: str, items: list[str]) -> None:
        doc.add_heading(label, level=2)
        if not items:
            doc.add_paragraph("未提供。")
            return
        for item in items:
            para = doc.add_paragraph(str(item), style="List Bullet")
            para.paragraph_format.space_after = Pt(2)

    @staticmethod
    def _set_cell_width(cell, width_inches: float) -> None:
        cell.width = Inches(width_inches)
        tc_pr = cell._tc.get_or_add_tcPr()
        tc_w = tc_pr.first_child_found_in("w:tcW")
        if tc_w is None:
            tc_w = OxmlElement("w:tcW")
            tc_pr.append(tc_w)
        tc_w.set(qn("w:w"), str(int(width_inches * 1440)))
        tc_w.set(qn("w:type"), "dxa")

    @staticmethod
    def _shade_cell(cell, fill: str) -> None:
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), fill)
        tc_pr.append(shd)
